import os
from docx import Document
import re
import pydot
from PIL import Image
import tkinter as tk
from tkinter.scrolledtext import ScrolledText


class WordUtility:
    @staticmethod
    def recognize_rule(word):
        return "URQ." in word

    @staticmethod
    def process_word(word):
        while word and word[0] != 'U':
            word = word[1:]
        pattern = r'\d$'
        while word and not re.search(pattern, word):
            word = word[:-1]
        return word

    @staticmethod
    def extract_name(cell):
        name = ""
        for paragraph in cell.paragraphs:
            name += paragraph.text + "\n"
        return name.strip()

    @staticmethod
    def extract_referenced_rules(cell, current_rule):
        referenced_rules = []
        for paragraph in cell.paragraphs:
            for word in paragraph.text.split():
                if WordUtility.recognize_rule(word):
                    processed_word = WordUtility.process_word(word)
                    if processed_word != current_rule:
                        referenced_rules.append(processed_word)
                    else:
                        referenced_rules.append("Error: self-reference")
        return referenced_rules


class Rule:
    def __init__(self, origin_document, signature, name, referenced_rules):
        self.origin_document = origin_document
        self.signature = signature
        self.name = name
        self.referenced_rules = referenced_rules

    def __str__(self):
        referenced_rules_str = ", ".join(str(rule) for rule in self.referenced_rules)
        return f"{self.signature}, {self.name}, ({referenced_rules_str})"


class RuleList:
    def __init__(self, rules):
        self.rules = rules

    def __str__(self):
        rules_str = [str(rule) for rule in self.rules]
        return '\n'.join(rules_str)

    def append(self, rule):
        self.rules.append(rule)

    def compile_referenced_rules(self):
        # iterate through each rule in the list of rules
        for rule in self.rules:
            # remove duplicates from the list of referenced rules and convert it back to a list
            rule.referenced_rules = list(set(rule.referenced_rules))
            # iterate through each referenced rule in the current rule's referenced rules
            for index, referenced_rule in enumerate(rule.referenced_rules):
                found = False
                # check if the referenced rule matches any existing rule's signature
                for existing_rule in self.rules:
                    if referenced_rule == existing_rule.signature:
                        # if a match is found, replace the referenced rule with the existing rule
                        rule.referenced_rules[index] = existing_rule
                        found = True
                        break
                # if no match is found and the referenced rule is not a self-reference error
                if not found and referenced_rule != "Error: self-reference":
                    # create a new rule with "Unknown" attributes and add it to the list of rules
                    new_rule = Rule("Unknown", referenced_rule, "Unknown", [])
                    self.rules.append(new_rule)
                    # replace the referenced rule with the newly created rule
                    rule.referenced_rules[index] = new_rule

    def specify(self, signature):
        current_rule = None
        found = False
        # iterate through each rule in the list of rules
        for rule in self.rules:
            # check if the signature of the current rule matches the specified signature
            if rule.signature == signature:
                current_rule = rule
                found = True
        # create an empty RuleList to store specified rules
        specified_rule_list = RuleList([])
        # if a rule with the specified signature is found
        if found:
            # add the current rule to the specified rule list
            specified_rule_list.rules.append(current_rule)
            # add referenced rules of the current rule to the specified rule list
            self.add_referenced_rules(current_rule, specified_rule_list)
        # return the specified rule list
        return specified_rule_list

    def add_referenced_rules(self, rule, specified_rule_list):
        # iterate through each referenced rule of the given rule
        for referenced_rule in rule.referenced_rules:
            # add the referenced rule to the specified rule list
            specified_rule_list.rules.append(referenced_rule)
            # recursively add referenced rules of the referenced rule to the specified rule list
            self.add_referenced_rules(referenced_rule, specified_rule_list)


class Graph:
    def __init__(self):
        self.graph = pydot.Dot(graph_type="graph", rankdir="UD", layout="circo")

    def create_graph(self, rule_list):
        # initialise a dictionary to store nodes of the graph
        nodes = {}
        # iterate through each rule in the list of rules
        for rule in rule_list.rules:
            # create a node for the rule and add it to the dictionary
            nodes[rule.signature] = pydot.Node(rule.signature,
                                               label=rule.signature + '\n' + rule.name + '\n' + rule.origin_document,
                                               shape="box")
            # add the node to the graph
            self.graph.add_node(nodes[rule.signature])
        # iterate through each rule in the list of rules
        for rule in rule_list.rules:
            # iterate through each referenced rule in the current rule's referenced rules
            for referenced_rule in rule.referenced_rules:
                # check if the referenced rule is an instance of Rule
                if isinstance(referenced_rule, Rule):
                    # create an edge from the current rule's node to the referenced rule's node
                    edge = pydot.Edge(nodes[rule.signature], nodes[referenced_rule.signature], dir="forward")
                    # add the edge to the graph
                    self.graph.add_edge(edge)
        # write the graph to an image file
        self.graph.write('graph.png', format='png', encoding='utf-8')
        # open and display the image
        img = Image.open('graph.png')
        img.show()


class Folder:
    def __init__(self, folder_path, rule_list):
        self.folder_path = folder_path
        self.rule_list = rule_list
        self.folder = os.path.abspath(self.folder_path)
        self.files = []
        self.error_message = []

    def check_validity(self):
        if not os.path.exists(self.folder) or not os.path.isdir(self.folder):
            self.error_message.append("This folder does not exist.")
        else:
            files = os.listdir(self.folder)
            if not files:
                self.error_message.append("This folder is empty.")
            else:
                files = [file for file in files if file.endswith((".doc", ".docx"))]
                if not files:
                    self.error_message.append("This folder does not contain any Microsoft Word documents.")
                else:
                    self.files = files

        return not bool(self.error_message)

    def process_files(self):
        # iterate through each file in the list of files
        for file in self.files:
            # construct the full file path
            file_path = os.path.join(self.folder, file)
            # extract the file name from the file path
            file_name = os.path.basename(file_path)
            # create a Document object from the file
            document = Document(file_path)
            # process tables in the document
            self.process_tables(document, file_name)

    def process_tables(self, document, file_name):
        # iterate through each table in the document
        for table in document.tables:
            # iterate through each row in the table
            for row_index, row in enumerate(table.rows):
                # ensure that the row index is within bounds
                if row_index < len(table.rows) - 1:
                    # check if the row has two cells
                    if len(row.cells) == 2:
                        # extract the cells
                        first_cell = row.cells[0]
                        second_cell = row.cells[1]
                        # get the first cell of the next row
                        third_cell = table.rows[row_index + 1].cells[0]
                        # check if the first cell is not empty
                        if first_cell.text.strip():
                            # extract the first word from the first cell's text
                            first_word = first_cell.paragraphs[0].text.split()[0]
                            # check if the first word is a signature
                            if WordUtility.recognize_rule(first_word):
                                # extract the name from the second cell
                                name = WordUtility.extract_name(second_cell)
                                # extract the referenced rules from the third cell
                                referenced_rules = WordUtility.extract_referenced_rules(third_cell, first_word)
                                # create a new rule object
                                new_rule = Rule(file_name, first_word, name, referenced_rules)
                                # add the new rule object to the list of rules
                                self.rule_list.append(new_rule)


class App:
    def __init__(self):
        self.window = None
        self.folder_label = None
        self.folder_entry = None
        self.create_graph = None
        self.rule_signature_label = None
        self.rule_signature_entry = None
        self.create_specified_graph = None
        self.output_text = None

    def create_gui(self):
        window = tk.Tk()
        window.title("Rule Processor")

        left_frame = tk.Frame(window)
        left_frame.grid(row=0, column=0, padx=10, pady=10)

        self.folder_label = tk.Label(left_frame, text="Enter Folder Path:")
        self.folder_label.grid(row=0, column=0, padx=10, pady=5)

        self.folder_entry = tk.Entry(left_frame, width=50)
        self.folder_entry.grid(row=1, column=0, padx=10, pady=5)

        self.create_graph = tk.Button(left_frame, text="Create Graph",
                                      command=self.create_graph_clicked, width=25)
        self.create_graph.grid(row=2, column=0, padx=10, pady=5)

        self.rule_signature_label = tk.Label(left_frame, text="Enter Rule Signature:")
        self.rule_signature_label.grid(row=3, column=0, padx=10, pady=5)

        self.rule_signature_entry = tk.Entry(left_frame, width=50)
        self.rule_signature_entry.grid(row=4, column=0, padx=10, pady=5)

        self.create_specified_graph = tk.Button(left_frame, text="Create Specified Graph",
                                                command=self.create_specified_graph_clicked, width=25)
        self.create_specified_graph.grid(row=5, column=0, padx=10, pady=5)

        right_frame = tk.Frame(window)
        right_frame.grid(row=0, column=1, padx=10, pady=10, sticky="nsew")

        self.output_text = ScrolledText(right_frame)
        self.output_text.pack(fill="both", expand=True)

        window.columnconfigure(1, weight=1)
        window.rowconfigure(0, weight=1)

        window.mainloop()

    def create_graph_clicked(self):
        folder_path = self.folder_entry.get()
        if not folder_path:  # Check if the folder path is not provided
            self.display_error_message("Please enter a folder path.")
        else:
            self.run(folder_path, False, None)

    def create_specified_graph_clicked(self):
        folder_path = self.folder_entry.get()
        rule_signature = self.rule_signature_entry.get()
        if not folder_path:  # Check if the folder path is not provided
            self.display_error_message("Please enter a folder path.")
        else:
            if rule_signature:  # If folder path is provided, proceed to create specified graph
                self.run(folder_path, True, rule_signature)
            else:
                self.display_error_message("Please enter a rule signature.")

    def run(self, folder_path, specify, rule_signature):
        # create an empty RuleList object, a Folder object and a Graph object
        rule_list = RuleList([])
        folder = Folder(folder_path, rule_list)
        graph = Graph()
        # check if the folder is valid
        if folder.check_validity():
            # process files in the folder
            folder.process_files()
            # compile referenced rules for all rules in the rule list
            rule_list.compile_referenced_rules()
            # if not creating a specified graph, display all rules and create a graph
            if not specify:
                self.display_rules(rule_list)
                graph.create_graph(rule_list)
            # if creating a specified graph
            else:
                # create a specified rule list
                new_rule_list = rule_list.specify(rule_signature)
                # if the specified rule list is not empty
                if new_rule_list.rules:
                    # display the specified rules and create a graph
                    self.display_rules(new_rule_list)
                    graph.create_graph(new_rule_list)
                # if the specified rule list is empty, display an error message
                else:
                    self.display_error_message("Invalid signature.")
        # if the folder is not valid, display an error message
        else:
            error_message = folder.error_message
            self.display_error_message('\n'.join(error_message))

    def display_rules(self, rule_list):
        rule_text = str(rule_list)
        self.output_text.delete(1.0, tk.END)
        self.output_text.insert(tk.END, rule_text)

    def display_error_message(self, text):
        self.output_text.delete(1.0, tk.END)
        self.output_text.insert(tk.END, text)


if __name__ == "__main__":
    app = App()
    app.create_gui()